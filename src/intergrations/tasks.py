# src/integrations/tasks.py
# ─────────────────────────────────────────────────────────────────────────────
# CoopData – Celery task definitions
# TOR §3.2 – Mambu API integration, KPI computation, report generation
# ─────────────────────────────────────────────────────────────────────────────
import logging
from datetime import datetime

from celery import shared_task
from django.conf import settings
from django.utils import timezone

logger = logging.getLogger(__name__)


# ═════════════════════════════════════════════════════════════════════════════
# KPI COMPUTATION
# ═════════════════════════════════════════════════════════════════════════════

@shared_task(bind=True, max_retries=3, default_retry_delay=60)
def compute_financial_kpis(self, summary_pk: int):
    """
    Called after a SACCOFinancialSummary is submitted.
    Runs compute_kpis() and then evaluates DataValidationRules.
    """
    from cooperatives.models import SACCOFinancialSummary
    try:
        summary = SACCOFinancialSummary.objects.get(pk=summary_pk)
        summary.compute_kpis()
        _run_validation_rules(summary)
        _compute_training_kpi(summary)
        logger.info("KPIs computed for summary #%s (%s)", summary_pk, summary.cooperative.name)
    except SACCOFinancialSummary.DoesNotExist:
        logger.error("SACCOFinancialSummary #%s not found.", summary_pk)
    except Exception as exc:
        logger.exception("Error computing KPIs for #%s", summary_pk)
        raise self.retry(exc=exc)


def _run_validation_rules(summary):
    """Evaluate active DataValidationRules against a financial summary."""
    from analytics.models import DataQualityAlert, DataValidationRule
    from simpleeval import simple_eval

    rules = DataValidationRule.objects.filter(
        is_active=True,
        applies_to_model="cooperatives.SACCOFinancialSummary",
    )
    for rule in rules:
        try:
            context = {
                f.name: getattr(summary, f.name)
                for f in summary._meta.fields
                if getattr(summary, f.name) is not None
            }
            violated = simple_eval(rule.rule_expression, names=context)
            if violated:
                msg = rule.error_message_template.format(
                    field=rule.applies_to_field or "N/A",
                    value=context.get(rule.applies_to_field, ""),
                )
                DataQualityAlert.objects.get_or_create(
                    rule=rule,
                    cooperative_id=summary.cooperative_id,
                    record_id=summary.pk,
                    is_resolved=False,
                    defaults={
                        "content_type_label": "SACCOFinancialSummary",
                        "field_name": rule.applies_to_field or "",
                        "message": msg or rule.name,
                    },
                )
        except Exception as exc:
            logger.warning("Rule '%s' evaluation failed: %s", rule.name, exc)


def _compute_training_kpi(summary):
    """Compute and store training hours per member KPI."""
    from cooperatives.models import TrainingRecord
    from django.db.models import Sum

    hours = (
        TrainingRecord.objects.filter(
            cooperative=summary.cooperative,
            training_date__gte=summary.period_start,
            training_date__lte=summary.period_end,
        ).aggregate(total=Sum("duration_hours"))["total"]
        or 0
    )
    if summary.total_members and summary.total_members > 0:
        summary.kpi_training_hours_per_member = round(
            float(hours) / summary.total_members, 2
        )
        summary.save(update_fields=["kpi_training_hours_per_member"])


# ═════════════════════════════════════════════════════════════════════════════
# MAMBU SYNC
# ═════════════════════════════════════════════════════════════════════════════

@shared_task(bind=True, max_retries=2, default_retry_delay=300)
def sync_mambu_cooperative(self, cooperative_id: int):
    """
    Pull latest data from Mambu for a cooperative.
    TOR §3.2 – API Integration with Mambu.
    """
    from integrations.mambu import MambuAPIError, sync_cooperative_from_mambu
    try:
        result = sync_cooperative_from_mambu(cooperative_id)
        if result["errors"]:
            logger.warning("Mambu sync partial for coop #%s: %s", cooperative_id, result["errors"])
        else:
            logger.info("Mambu sync complete for coop #%s: %s", cooperative_id, result)
        return result
    except MambuAPIError as exc:
        logger.error("Mambu API error for coop #%s: %s", cooperative_id, exc)
        raise self.retry(exc=exc)
    except Exception as exc:
        logger.exception("Unexpected Mambu sync error for coop #%s", cooperative_id)
        raise self.retry(exc=exc)


@shared_task
def sync_all_mambu_cooperatives():
    """
    Periodic task: sync all cooperatives that have a Mambu encoded key.
    Scheduled via Celery Beat (configurable interval in .env).
    """
    from cooperatives.models import Cooperative
    coops = Cooperative.objects.filter(
        mambu_encoded_key__isnull=False,
        status="active",
    ).values_list("id", flat=True)
    for coop_id in coops:
        sync_mambu_cooperative.delay(coop_id)
    logger.info("Queued Mambu sync for %d cooperatives.", len(coops))


# ═════════════════════════════════════════════════════════════════════════════
# REPORT / EXPORT GENERATION
# ═════════════════════════════════════════════════════════════════════════════

@shared_task(bind=True, max_retries=2)
def generate_export(self, job_pk: int):
    """
    Generate a report file (PDF / Excel / Word) for an ExportJob.
    TOR §3.4 – exportable reports PDF, Excel, Word.
    """
    from analytics.models import ExportJob
    try:
        job = ExportJob.objects.get(pk=job_pk)
        job.status = ExportJob.STATUS_RUNNING
        job.save(update_fields=["status"])

        if job.format == "xlsx":
            url = _generate_excel(job)
        elif job.format == "pdf":
            url = _generate_pdf(job)
        elif job.format == "docx":
            url = _generate_word(job)
        else:
            raise ValueError(f"Unknown format: {job.format}")

        job.file_url = url
        job.status = ExportJob.STATUS_DONE
        job.completed_at = timezone.now()
        job.save(update_fields=["file_url", "status", "completed_at"])
        logger.info("Export #%s complete: %s", job_pk, url)
    except ExportJob.DoesNotExist:
        logger.error("ExportJob #%s not found.", job_pk)
    except Exception as exc:
        logger.exception("Export #%s failed", job_pk)
        try:
            job.status = ExportJob.STATUS_FAILED
            job.error_message = str(exc)
            job.save(update_fields=["status", "error_message"])
        except Exception:
            pass
        raise self.retry(exc=exc)


def _generate_excel(job) -> str:
    """Generate Excel report using openpyxl and upload to Cloudinary."""
    import io
    import openpyxl
    from openpyxl.styles import Font, PatternFill
    from django.core.files.base import ContentFile
    import cloudinary.uploader

    from cooperatives.models import SACCOFinancialSummary

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "KPI Report"

    # Header row
    headers = [
        "Cooperative", "Period", "Total Members", "Female Members", "Youth Members",
        "Gross Loan Portfolio (SZL)", "PAR 30 (%)", "Liquidity Ratio (%)",
        "Capital Adequacy (%)", "ROA (%)", "OSS (%)",
    ]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1B4F72")

    # Data rows
    params = job.parameters_snapshot or {}
    coop_ids = params.get("cooperative_ids")
    qs = SACCOFinancialSummary.objects.filter(is_verified=True).select_related("cooperative")
    if coop_ids:
        qs = qs.filter(cooperative_id__in=coop_ids)
    qs = qs.order_by("cooperative__name", "-period_end")

    for row_idx, s in enumerate(qs, 2):
        ws.cell(row=row_idx, column=1, value=s.cooperative.name)
        ws.cell(row=row_idx, column=2, value=f"{s.period_start} – {s.period_end}")
        ws.cell(row=row_idx, column=3, value=s.total_members)
        ws.cell(row=row_idx, column=4, value=s.female_members)
        ws.cell(row=row_idx, column=5, value=s.youth_members)
        ws.cell(row=row_idx, column=6, value=float(s.gross_loan_portfolio or 0))
        ws.cell(row=row_idx, column=7, value=round(float(s.kpi_delinquency_rate or 0) * 100, 2))
        ws.cell(row=row_idx, column=8, value=round(float(s.kpi_liquidity_ratio or 0) * 100, 2))
        ws.cell(row=row_idx, column=9, value=round(float(s.kpi_capital_adequacy or 0) * 100, 2))
        ws.cell(row=row_idx, column=10, value=round(float(s.kpi_roa or 0) * 100, 2))
        ws.cell(row=row_idx, column=11, value=round(float(s.kpi_operational_self_sufficiency or 0) * 100, 2))

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    result = cloudinary.uploader.upload(
        buf,
        resource_type="raw",
        folder="coopdata/exports",
        public_id=f"kpi_report_{job.pk}",
        format="xlsx",
    )
    return result["secure_url"]


def _generate_pdf(job) -> str:
    """Generate PDF using ReportLab and upload to Cloudinary."""
    import io
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    import cloudinary.uploader

    from cooperatives.models import SACCOFinancialSummary

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4))
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("CoopData – KPI Report", styles["Title"]))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["Normal"]))
    elements.append(Spacer(1, 12))

    qs = SACCOFinancialSummary.objects.filter(is_verified=True).select_related("cooperative").order_by("cooperative__name", "-period_end")
    params = job.parameters_snapshot or {}
    if params.get("cooperative_ids"):
        qs = qs.filter(cooperative_id__in=params["cooperative_ids"])

    table_data = [["Cooperative", "Period", "Members", "PAR30%", "LIQ%", "CAP%", "ROA%"]]
    for s in qs:
        table_data.append([
            s.cooperative.name[:30],
            str(s.period_end),
            str(s.total_members),
            f"{float(s.kpi_delinquency_rate or 0)*100:.1f}%",
            f"{float(s.kpi_liquidity_ratio or 0)*100:.1f}%",
            f"{float(s.kpi_capital_adequacy or 0)*100:.1f}%",
            f"{float(s.kpi_roa or 0)*100:.1f}%",
        ])

    t = Table(table_data)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1B4F72")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EBF5FB")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
    ]))
    elements.append(t)
    doc.build(elements)
    buf.seek(0)
    result = cloudinary.uploader.upload(
        buf, resource_type="raw", folder="coopdata/exports",
        public_id=f"kpi_report_{job.pk}", format="pdf",
    )
    return result["secure_url"]


def _generate_word(job) -> str:
    """Generate Word .docx report using python-docx and upload to Cloudinary."""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    import cloudinary.uploader

    from cooperatives.models import SACCOFinancialSummary

    doc = Document()
    doc.add_heading("CoopData – KPI Report", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    qs = SACCOFinancialSummary.objects.filter(is_verified=True).select_related("cooperative").order_by("cooperative__name", "-period_end")
    params = job.parameters_snapshot or {}
    if params.get("cooperative_ids"):
        qs = qs.filter(cooperative_id__in=params["cooperative_ids"])

    headers = ["Cooperative", "Period", "Members", "PAR30%", "LIQ%", "CAP%", "ROA%"]
    rows = len(list(qs))
    table = doc.add_table(rows=rows + 1, cols=len(headers))
    table.style = "Table Grid"

    for col_i, h in enumerate(headers):
        cell = table.rows[0].cells[col_i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for row_i, s in enumerate(qs, 1):
        row = table.rows[row_i].cells
        row[0].text = s.cooperative.name
        row[1].text = str(s.period_end)
        row[2].text = str(s.total_members)
        row[3].text = f"{float(s.kpi_delinquency_rate or 0)*100:.1f}%"
        row[4].text = f"{float(s.kpi_liquidity_ratio or 0)*100:.1f}%"
        row[5].text = f"{float(s.kpi_capital_adequacy or 0)*100:.1f}%"
        row[6].text = f"{float(s.kpi_roa or 0)*100:.1f}%"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    result = cloudinary.uploader.upload(
        buf, resource_type="raw", folder="coopdata/exports",
        public_id=f"kpi_report_{job.pk}", format="docx",
    )
    return result["secure_url"]


# ═════════════════════════════════════════════════════════════════════════════
# CELERY BEAT SCHEDULE (put in settings.py / celery.py)
# ═════════════════════════════════════════════════════════════════════════════
# CELERY_BEAT_SCHEDULE = {
#     "sync-mambu-every-hour": {
#         "task": "integrations.tasks.sync_all_mambu_cooperatives",
#         "schedule": crontab(minute=0),   # top of every hour
#     },
# }