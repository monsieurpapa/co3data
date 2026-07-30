"""CO3DATA — Celery Background Tasks"""
import logging
from datetime import date, datetime, timedelta

from celery import shared_task
from django.utils import timezone

logger = logging.getLogger("co3data.tasks")


@shared_task(bind=True, max_retries=3, default_retry_delay=60)
def compute_kpis_for_cooperative(self, cooperative_id: int, period_start: str, period_end: str):
    from cooperatives.models import Cooperative, Member, ProductionRecord, FinancialRecord
    try:
        cooperative = Cooperative.objects.get(pk=cooperative_id)
        start, end = date.fromisoformat(period_start), date.fromisoformat(period_end)
        members = Member.objects.filter(cooperative=cooperative, is_active=True)

        production = ProductionRecord.objects.filter(
            member__cooperative=cooperative,
        ).filter(
            harvest_date__range=(start, end),
        ) | ProductionRecord.objects.filter(
            member__cooperative=cooperative,
            purchase_date__range=(start, end),
        )
        total_production_kg = sum((p.quantity_kg or 0) for p in production.distinct())

        financial = FinancialRecord.objects.filter(cooperative=cooperative, transaction_date__range=(start, end))
        income = sum((f.amount for f in financial.filter(transaction_type="income")), start=0)
        expenses = sum((f.amount for f in financial.filter(transaction_type="expense")), start=0)

        total_m = members.count()
        female_m = members.filter(gender="female").count()
        youth_m = members.filter(age_group="youth").count()

        result = {
            "cooperative": cooperative.name,
            "total_members": total_m,
            "female_members": female_m,
            "youth_members": youth_m,
            "female_participation_rate": round(female_m / total_m * 100, 2) if total_m else 0,
            "youth_participation_rate": round(youth_m / total_m * 100, 2) if total_m else 0,
            "total_production_kg": float(total_production_kg),
            "production_per_member": round(float(total_production_kg) / total_m, 2) if total_m else 0,
            "income": float(income),
            "expenses": float(expenses),
            "net_income": float(income - expenses),
        }
        logger.info(f"Computed KPIs for {cooperative.name} [{start}–{end}]")
        return result
    except Exception as exc:
        raise self.retry(exc=exc)


@shared_task
def compute_all_cooperative_kpis():
    from cooperatives.models import Cooperative
    today = timezone.now().date()
    period_end = today.replace(day=1) - timedelta(days=1)
    period_start = period_end.replace(day=1)
    for coop in Cooperative.objects.all():
        compute_kpis_for_cooperative.delay(coop.pk, period_start.isoformat(), period_end.isoformat())


@shared_task(bind=True, max_retries=3)
def run_data_quality_checks(self, cooperative_id: int = None):
    from analytics.services import ValidationService
    from cooperatives.models import Cooperative, ProductionRecord, FinancialRecord

    coops = Cooperative.objects.all()
    if cooperative_id:
        coops = coops.filter(pk=cooperative_id)
    coop_ids = coops.values_list("id", flat=True)

    checked = 0
    for record in ProductionRecord.objects.filter(member__cooperative_id__in=coop_ids):
        ValidationService.validate_record(record)
        checked += 1
    for record in FinancialRecord.objects.filter(cooperative_id__in=coop_ids):
        ValidationService.validate_record(record)
        checked += 1

    logger.info(f"Data quality check completed: {checked} records checked")
    return {"records_checked": checked}


@shared_task
def cleanup_old_sync_logs(days_to_keep: int = 90):
    from sync.models import SyncLog
    cutoff = timezone.now() - timedelta(days=days_to_keep)
    deleted, _ = SyncLog.objects.filter(sync_start_time__lt=cutoff).delete()
    logger.info(f"Deleted {deleted} old sync logs")
    return {"deleted": deleted}


@shared_task(bind=True, max_retries=3)
def generate_export(self, job_pk: int):
    """Generate a report file (PDF / Excel / Word) for an analytics.ExportJob."""
    from analytics.models import ExportJob
    try:
        job = ExportJob.objects.get(pk=job_pk)
        job.status = ExportJob.STATUS_RUNNING
        job.save(update_fields=["status"])

        if job.format == "xlsx":
            url = _generate_excel(job)
        elif job.format == "pdf":
            url = _generate_pdf(job)
        elif job.format == "docx":
            url = _generate_word(job)
        else:
            raise ValueError(f"Unknown format: {job.format}")

        job.file_url = url
        job.status = ExportJob.STATUS_DONE
        job.completed_at = timezone.now()
        job.save(update_fields=["file_url", "status", "completed_at"])
        logger.info("Export #%s complete: %s", job_pk, url)
    except ExportJob.DoesNotExist:
        logger.error("ExportJob #%s not found.", job_pk)
    except Exception as exc:
        logger.exception("Export #%s failed", job_pk)
        try:
            job.status = ExportJob.STATUS_FAILED
            job.error_message = str(exc)
            job.save(update_fields=["status", "error_message"])
        except Exception:
            pass
        raise self.retry(exc=exc)


def _cooperative_summary_rows(job):
    """Per-cooperative production/sales summary rows for the export report."""
    from cooperatives.models import Cooperative, CooperativeSale
    from django.db.models import Sum, Count, Q

    params = job.parameters_snapshot or {}
    coops = Cooperative.objects.all()
    if params.get("cooperative_ids"):
        coops = coops.filter(id__in=params["cooperative_ids"])

    coops = coops.annotate(
        total_members=Count("members", distinct=True),
        female_members=Count("members", filter=Q(members__gender="female"), distinct=True),
        youth_members=Count("members", filter=Q(members__age_group="youth"), distinct=True),
        total_production_kg=Sum("members__production_records__quantity_kg"),
        total_sales_value=Sum("sales__total_value"),
    ).order_by("name")
    return list(coops)


def _generate_excel(job) -> str:
    """Generate an Excel cooperative summary report using openpyxl and upload to Cloudinary."""
    import io
    import openpyxl
    from openpyxl.styles import Font, PatternFill
    import cloudinary.uploader

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Rapport Coopératives"

    headers = ["Coopérative", "Membres", "Femmes", "Jeunes", "Production (kg)", "Ventes (valeur)"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1B4F72")

    for row_idx, coop in enumerate(_cooperative_summary_rows(job), 2):
        ws.cell(row=row_idx, column=1, value=coop.name)
        ws.cell(row=row_idx, column=2, value=coop.total_members)
        ws.cell(row=row_idx, column=3, value=coop.female_members)
        ws.cell(row=row_idx, column=4, value=coop.youth_members)
        ws.cell(row=row_idx, column=5, value=float(coop.total_production_kg or 0))
        ws.cell(row=row_idx, column=6, value=float(coop.total_sales_value or 0))

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    result = cloudinary.uploader.upload(
        buf, resource_type="raw", folder="co3data/exports",
        public_id=f"report_{job.pk}", format="xlsx",
    )
    return result["secure_url"]


def _generate_pdf(job) -> str:
    """Generate a PDF cooperative summary report using ReportLab and upload to Cloudinary."""
    import io
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    import cloudinary.uploader

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4))
    styles = getSampleStyleSheet()
    elements = [
        Paragraph("CO3DATA – Rapport des Coopératives", styles["Title"]),
        Paragraph(f"Généré : {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["Normal"]),
        Spacer(1, 12),
    ]

    table_data = [["Coopérative", "Membres", "Femmes", "Jeunes", "Production (kg)", "Ventes"]]
    for coop in _cooperative_summary_rows(job):
        table_data.append([
            coop.name[:30],
            str(coop.total_members),
            str(coop.female_members),
            str(coop.youth_members),
            f"{float(coop.total_production_kg or 0):.1f}",
            f"{float(coop.total_sales_value or 0):.2f}",
        ])

    t = Table(table_data)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1B4F72")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EBF5FB")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
    ]))
    elements.append(t)
    doc.build(elements)
    buf.seek(0)
    result = cloudinary.uploader.upload(
        buf, resource_type="raw", folder="co3data/exports",
        public_id=f"report_{job.pk}", format="pdf",
    )
    return result["secure_url"]


def _generate_word(job) -> str:
    """Generate a Word .docx cooperative summary report using python-docx and upload to Cloudinary."""
    import io
    from docx import Document
    from docx.shared import Pt
    import cloudinary.uploader

    doc = Document()
    doc.add_heading("CO3DATA – Rapport des Coopératives", level=1)
    doc.add_paragraph(f"Généré : {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Coopérative", "Membres", "Femmes", "Jeunes", "Production (kg)", "Ventes"]):
        hdr[i].text = h

    for coop in _cooperative_summary_rows(job):
        row = table.add_row().cells
        row[0].text = coop.name
        row[1].text = str(coop.total_members)
        row[2].text = str(coop.female_members)
        row[3].text = str(coop.youth_members)
        row[4].text = f"{float(coop.total_production_kg or 0):.1f}"
        row[5].text = f"{float(coop.total_sales_value or 0):.2f}"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    result = cloudinary.uploader.upload(
        buf, resource_type="raw", folder="co3data/exports",
        public_id=f"report_{job.pk}", format="docx",
    )
    return result["secure_url"]
